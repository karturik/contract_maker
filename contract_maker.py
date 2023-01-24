from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import random
import pandas as pd
from docx2pdf import convert

read_file = pd.read_excel ('users_data.xlsx', sheet_name='Ответы на форму (1)')
read_file.to_csv ('users_data.csv', index = False, header=True, sep='^')

smena_padezha = {'Приднестровская Молдавская Республика': 'Приднестровской Молдавской Республики', 'Республика Беларусь':'Республики Беларусь', 'Российская Федерация':'Российской Федерации', 'Украина':'Украины', 'Республика Молдова':'Республики Молдовы', 'Грузия':'Грузии'}

with open('users_data.csv', 'r', encoding='utf-8') as file_src:
    users_data = file_src.read().strip().split('\n')[1:]
    file_src.close()

    # print(users_data)
for row in users_data:
    row = row.replace('""', '"').split('^')
    print(row)
    doc_number = '11/22/3333344555'
    doc_start_date = '21.32.3221'
    user_name = row[3].strip()
    user_otchestvo = row[4].strip()
    user_surname = row[2].strip()
    user_fio = f'{user_name} {user_otchestvo} {user_surname}'
    user_country = row[17]
    try:
        country_v_rod_padezhe = smena_padezha[user_country]
    except:
        country_v_rod_padezhe = user_country
    user_oblast = row[18]
    user_city = row[19]
    user_zip = row[20]
    user_street = row[21]
    user_adress = f'{user_country}, {user_zip}, {user_oblast}, {user_city}, {user_street}'
    user_date_birth = row[5].split(" ")[0]
    pass_number = row[11]
    pass_date = row[13].split(" ")[0]
    pass_organ_vidachy = row[15]
    r_schet_number = row[24]
    k_schet_number = '12345678910'
    bank_name = row[22]
    bank_adress = row[23]
    bank_bic = row[31]
    user_mail = row[25]

    ip_label = 'Индивидуальный предприниматель,'
    grazhdanin_label = 'гражданин'
    svidetelstvo_ip_label = 'Свидетельства о государственной регистрации индивидуального предпринимателя '
    nomer_svidetelstva = ' ' + '1234567891011'
    fizicheskoe_lico_label = ' , действующий как физическое лицо'
    registration_data_label = 'зарегистрированный в ФНС в качестве налогоплательщика налога на профессиональный доход в соответствии с Федеральным законом от 27 ноября 2018 года No. 422-ФЗ, действующий на основании '

    user_status = 'гражданин'

    if user_status == 'гражданин':
        ip_or_grazhdanin_label = grazhdanin_label
        svidetelstvo_ip_label = ''
        fizicheskoe_lico_label = fizicheskoe_lico_label
        nomer_svidetelstva = ''
        registration_data_label = ''
    else:
        ip_or_grazhdanin_label = ip_label
        nomer_svidetelstva = ' No.' + nomer_svidetelstva
        fizicheskoe_lico_label = ''

    user_status = f'{ip_or_grazhdanin_label} ' \
                  f'{registration_data_label}' \
                  f'{svidetelstvo_ip_label}' \
                  f'{country_v_rod_padezhe}{nomer_svidetelstva}, {user_surname} {user_name} {user_otchestvo}'\
                  f', именуемый в дальнейшем «Исполнитель»{fizicheskoe_lico_label}' \
                  f' с другой стороны, совместно именуемые «Стороны»,' \
                  f' а индивидуально — «Сторона», заключили настоящий Рамочный договор на выполнение Работ, оказание Услуг (далее — Договор)' \
                  f' на изложенных ниже условиях'

    document = Document('doc_project.docx')

    paragraph0 = document.paragraphs[0]
    paragraph0.text = paragraph0.text.replace('doc_number', doc_number)
    for run in paragraph0.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True

    paragraph7 = document.paragraphs[7]
    paragraph7.text = user_status
    for run in paragraph7.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)

    paragraph107 = document.paragraphs[107]
    paragraph107.text = paragraph107.text.replace('user_mail', user_mail)
    # paragraph7.text = user_status
    for run in paragraph107.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        # print(run.text)

    num = 0
    doc_info = document.tables[0].rows[0].cells[1]
    doc_info.text = doc_info.text.replace('doc_start_date', doc_start_date)
    for paragraph in doc_info.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    end_data = document.tables[1].rows[0].cells[2]
    num = 0
    for paragraph in end_data.paragraphs:
        # print(num, paragraph.text)
        num += 1

    # ФИО
    end_paragraph2 = end_data.paragraphs[2]
    end_paragraph2.text = user_fio
    for run in end_paragraph2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    end_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Адрес прописки
    end_paragraph4 = end_data.paragraphs[4]
    end_paragraph4.text = user_adress
    for run in end_paragraph4.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    end_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Дата рождения
    end_paragraph6 = end_data.paragraphs[6]
    end_paragraph6.text = f'{user_date_birth} года рождения'
    for run in end_paragraph6.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    end_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Паспорт
    end_paragraph7 = end_data.paragraphs[7]
    end_paragraph7.text = f'Паспорт {pass_number} выдан\n{pass_date}, {pass_organ_vidachy}'
    for run in end_paragraph7.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    end_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # расчетный счет
    end_paragraphs9_12 = end_data.paragraphs[9:14]
    end_paragraphs9_12[0].text = f'Р/ счет No. {r_schet_number}'
    end_paragraphs9_12[1].text = f'К/ счет No. {k_schet_number}'
    end_paragraphs9_12[2].text = f'{bank_name}'
    end_paragraphs9_12[3].text = f'{bank_adress}'
    end_paragraphs9_12[4].text = f'{bank_bic}'
    for paragraph in end_paragraphs9_12:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


    end_paragraphs19_24 = end_data.paragraphs[19:25]
    end_paragraphs19_24[0].text = f'{user_fio}'
    end_paragraphs19_24[-1].text = f'_________________ {user_name[0].capitalize()}.{user_otchestvo[0].capitalize()}.{user_surname}'
    for paragraph in end_paragraphs19_24:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    file_name = random.randint(10000, 100000)
    try:
        path = f'docs/{user_name[0].capitalize()}.{user_otchestvo[0].capitalize()}.{user_surname}/doc N {file_name}'
        # os.mkdir(f'docs/{user_name[0].capitalize()}.{user_otchestvo[0].capitalize()}.{user_surname}')
        os.makedirs(path)
        document.save(f'{path}/{user_fio}.docx')
        convert(f'{path}/{user_fio}.docx', f'{path}/{user_fio}.pdf')
    except Exception as e:
        print(e)


